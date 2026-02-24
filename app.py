from flask import Flask
from flask_cors import CORS
from config import Config
from extensions import db

def create_app():
    app = Flask(__name__)
    app.config.from_object(Config)
    
    CORS(app, resources={r"/api/*": {"origins": ["http://localhost:5173", "https://iuristatech.com", "https://www.iuristatech.com", "http://173.212.225.148", "https://n8n-bot-back-cristian.gnuu1e.easypanel.host"]}})
    db.init_app(app)

    # File upload config
    import os
    app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max
    UPLOAD_FOLDER = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'uploads')
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

    from admin_routes import admin_bp
    app.register_blueprint(admin_bp)

    @app.route('/')
    def index():
         return "LegalTech Backend API is running. Use /api/health or /api/login."
    
    @app.route('/api/health', methods=['GET'])
    def health_check():
        return {"status": "ok", "message": "Backend is running"}

    
    @app.route('/api/login', methods=['POST'])
    def login():
        from flask import request
        from models import User
        
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        
        # In a real app, hash password and compare
        user = User.query.filter_by(email=email).first()
        
        if user and user.password_hash == password:
             if not user.is_approved:
                 return {"status": "error", "message": "Tu cuenta está pendiente de aprobación por un administrador."}, 403

             return {
                 "status": "success", 
                 "token": "dummy-jwt-token", 
                 "user": {
                     "id": user.id, 
                     "email": user.email, 
                     "name": user.full_name,
                     "role": user.role,
                     "is_admin": user.is_admin
                 }
             }, 200
            
        return {"status": "error", "message": "Credenciales inválidas"}, 401

    @app.route('/api/register', methods=['POST'])
    def register():
        from flask import request
        from models import User
        
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        name = data.get('name')
        
        if not email or not password:
            return {"error": "Email y contraseña requeridos"}, 400
            
        if User.query.filter_by(email=email).first():
            return {"error": "El correo ya está registrado"}, 400
            
        # Determine if this is the first user (make admin/approved automatically)
        user_count = User.query.count()
        is_first_user = (user_count == 0)
        
        new_user = User(
            email=email,
            password_hash=password, # TODO: Hash
            full_name=name,
            role='admin' if is_first_user else 'client',
            is_admin=is_first_user,
            is_approved=is_first_user # Auto-approve first user
        )
        
        db.session.add(new_user)
        db.session.commit()
        
        msg = "Usuario creado exitosamente."
        if not is_first_user:
            msg += " Su cuenta está pendiente de aprobación."
            
        return {"status": "success", "message": msg, "userId": new_user.id}, 201

    @app.route('/api/upload', methods=['POST'])
    def upload_document():
        from flask import request
        import io
        
        if 'file' not in request.files:
            return {"error": "No se envió ningún archivo"}, 400
        
        file = request.files['file']
        if file.filename == '':
            return {"error": "Nombre de archivo vacío"}, 400
        
        filename = file.filename.lower()
        allowed_extensions = {'.pdf', '.docx', '.txt'}
        ext = '.' + filename.rsplit('.', 1)[-1] if '.' in filename else ''
        
        if ext not in allowed_extensions:
            return {"error": f"Tipo de archivo no soportado. Use: {', '.join(allowed_extensions)}"}, 400
        
        try:
            extracted_text = ''
            
            if ext == '.pdf':
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(file.read()))
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + '\n'
            
            elif ext == '.docx':
                from docx import Document
                doc = Document(io.BytesIO(file.read()))
                for para in doc.paragraphs:
                    if para.text.strip():
                        extracted_text += para.text + '\n'
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        extracted_text += row_text + '\n'
            
            elif ext == '.txt':
                content = file.read()
                # Try UTF-8 first, then latin-1
                try:
                    extracted_text = content.decode('utf-8')
                except UnicodeDecodeError:
                    extracted_text = content.decode('latin-1')
            
            if not extracted_text.strip():
                return {"error": "No se pudo extraer texto del archivo. El archivo puede estar vacío o protegido."}, 400
            
            return {
                "status": "success",
                "filename": file.filename,
                "text": extracted_text,
                "characters": len(extracted_text)
            }
        
        except Exception as e:
            print(f"Error processing file: {e}")
            return {"error": f"Error al procesar el archivo: {str(e)}"}, 500

    @app.route('/api/chat', methods=['POST'])
    def chat():
        from flask import request
        from models import Conversation, Message
        from services.chat_engine import generate_response
        
        data = request.get_json()
        user_id = data.get('userId') # Optional if anonymous
        message_text = data.get('message')
        conversation_id = data.get('conversationId')
        document_context = data.get('documentContext')  # {filename, text}
        
        if not message_text:
            return {"error": "Message required"}, 400
            
        # 1. Get or Create Conversation
        if not conversation_id:
            title = message_text[:40] + "..." if len(message_text) > 40 else message_text
            conv = Conversation(user_id=user_id, title=title)
            db.session.add(conv)
            db.session.commit()
            conversation_id = conv.id
        else:
            conv = db.session.get(Conversation, conversation_id)
            if not conv:
                return {"error": "Conversation not found"}, 404
        
        # 2. Save User Message (include document info if present)
        msg_content = message_text
        if document_context:
            msg_content = f"📎 [{document_context.get('filename', 'Documento')}]\n{message_text}"
        user_msg = Message(conversation_id=conversation_id, sender_role='user', content=msg_content)
        db.session.add(user_msg)
        
        # 3. Generate AI Response (with document context if present)
        ai_result = generate_response(message_text, conversation_id, document_context=document_context)
        ai_text = ai_result['text']
        
        # 4. Save AI Message
        ai_msg = Message(conversation_id=conversation_id, sender_role='assistant', content=ai_text)
        db.session.add(ai_msg)
        
        db.session.commit()
        
        
        return {
            "conversationId": conversation_id,
            "response": ai_text,
            "status": ai_result['status'],
            "suggestedActions": ai_result.get('suggested_actions', []),
            "title": conv.title
        }

    @app.route('/api/conversations/<int:user_id>', methods=['GET'])
    def get_conversations(user_id):
        from models import Conversation
        convs = Conversation.query.filter_by(user_id=user_id).order_by(Conversation.updated_at.desc()).all()
        
        result = []
        for c in convs:
            # Get last message / simple summary
            result.append({
                "id": c.id,
                "title": c.title,
                "updated_at": c.updated_at.isoformat()
            })
        return {"conversations": result}
        
    @app.route('/api/conversations/<int:conversation_id>/messages', methods=['GET'])
    def get_messages(conversation_id):
        from models import Message
        msgs = Message.query.filter_by(conversation_id=conversation_id).order_by(Message.created_at.asc()).all()
        
        result = []
        for m in msgs:
            result.append({
                "id": m.id,
                "role": m.sender_role,
                "content": m.content,
                "timestamp": m.created_at.isoformat()
            })
        return {"messages": result}

    @app.route('/api/user/<int:user_id>/profile', methods=['GET'])
    def get_user_profile(user_id):
        from models import User, Conversation, Message
        
        user = db.session.get(User, user_id)
        if not user:
            return {"error": "Usuario no encontrado"}, 404
        
        # Get stats
        total_conversations = Conversation.query.filter_by(user_id=user_id).count()
        total_messages = Message.query.join(Conversation).filter(Conversation.user_id == user_id).count()
        
        # Get recent conversations
        convs = Conversation.query.filter_by(user_id=user_id).order_by(Conversation.updated_at.desc()).limit(10).all()
        conversations = [{
            "id": c.id,
            "title": c.title,
            "status": c.status,
            "updated_at": c.updated_at.isoformat()
        } for c in convs]
        
        return {
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.full_name,
                "role": user.role,
                "is_admin": user.is_admin,
                "created_at": user.created_at.isoformat() if user.created_at else None
            },
            "stats": {
                "totalConversations": total_conversations,
                "totalMessages": total_messages
            },
            "conversations": conversations
        }

    @app.route('/api/user/<int:user_id>/profile', methods=['PUT'])
    def update_user_profile(user_id):
        from flask import request
        from models import User
        
        user = db.session.get(User, user_id)
        if not user:
            return {"error": "Usuario no encontrado"}, 404
        
        data = request.get_json()
        if 'name' in data:
            user.full_name = data['name']
        
        db.session.commit()
        
        return {
            "status": "success",
            "message": "Perfil actualizado",
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.full_name,
                "role": user.role,
                "is_admin": user.is_admin
            }
        }

        
    return app

if __name__ == '__main__':
    app = create_app()
    app.run(debug=True, port=5000)
